from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


DOCX_PATH = Path("/Users/joss66/trabajo/SGA/Guion_de_Entrevista-Habilidades_respondido.docx")


ANSWERS = {
    "¿Cuáles son los principales desafíos que enfrenta un líder para gestionar una organización en un entorno de turbulencia empresarial?": (
        "Uno de los principales desafíos es tomar decisiones con información incompleta y al mismo tiempo mantener al equipo enfocado. "
        "Como desarrollador, líder de equipo y cofundador, he aprendido que la turbulencia obliga a priorizar con mucha claridad: qué problema resolvemos, "
        "qué impacto tendrá y qué riesgos estamos aceptando. También exige comunicar el contexto, no solo las tareas. Cuando el equipo entiende por qué se cambia "
        "una prioridad o por qué se adopta una nueva tecnología, responde con más criterio y autonomía."
    ),
    "“Para empezar, ¿cómo definiría usted el liderazgo en el contexto de la gestión empresarial?”": (
        "Para mí, liderar es crear las condiciones para que un equipo pueda tomar buenas decisiones, ejecutar con calidad y crecer en el proceso. "
        "No lo veo como controlar cada movimiento, sino como alinear visión, responsabilidades y comunicación. En una empresa tecnológica, el liderazgo también implica "
        "traducir la estrategia del negocio a decisiones técnicas concretas, y traducir las capacidades técnicas del equipo en valor real para clientes y usuarios."
    ),
    "“¿Cuáles considera que son los principales retos a los que se enfrenta un líder hoy en día?”": (
        "El principal reto es mantener claridad en medio de demasiada información, herramientas y cambios de mercado. Hoy un líder debe aprender rápido, filtrar ruido "
        "y ayudar al equipo a enfocarse. Otro reto es incorporar tecnologías como la inteligencia artificial sin generar miedo. La IA no debe presentarse como un reemplazo "
        "de las personas, sino como una capa de apoyo que reduce tareas repetitivas, mejora la documentación, acelera análisis y permite que el talento humano se concentre "
        "en criterio, creatividad, relación con clientes y solución de problemas complejos."
    ),
    "“En su opinión, ¿Por qué es  importante la\xa0comunicación efectiva\xa0para un líder o empresario?”": (
        "La comunicación efectiva es importante porque evita que el equipo desperdicie energía interpretando supuestos. En desarrollo de software, una instrucción ambigua "
        "puede convertirse en deuda técnica, retrasos o frustración. Como empresario, también he visto que comunicar bien permite conectar tecnología, ventas, operación y clientes. "
        "La IA puede reforzar esta comunicación: ayuda a resumir reuniones, convertir conversaciones en tareas, documentar decisiones y detectar inconsistencias, pero la intención, "
        "la empatía y el criterio siguen siendo responsabilidad del líder."
    ),
    "“¿Podría compartir un ejemplo en el que la comunicación clara haya sido fundamental para la resolución de un problema o la consecución de un objetivo en su equipo?”": (
        "Un ejemplo frecuente ocurre cuando un producto cambia de alcance durante el desarrollo. En lugar de simplemente reasignar tareas, suelo reunir al equipo para explicar "
        "qué cambió, por qué cambió y cuál es el nuevo objetivo de negocio. Después convertimos esa conversación en criterios de aceptación, responsables y fechas claras. "
        "Cuando usamos herramientas de IA para resumir la reunión o generar una primera versión de documentación técnica, el equipo gana velocidad, pero lo esencial es que todos "
        "compartan la misma lectura del problema."
    ),
    "Segú su experiencia “En la toma de decisiones cuáles son los aspectos que debe tener en cuenta para disminuir el riesgo de incertidumbre? .": (
        "Para reducir la incertidumbre considero cuatro aspectos: datos disponibles, impacto esperado, costo de equivocarse y capacidad de reversar la decisión. "
        "En tecnología no todas las decisiones tienen el mismo peso: elegir una arquitectura, contratar una persona clave o cambiar una prioridad de producto requiere más análisis "
        "que una optimización puntual. Uso la IA como apoyo para comparar alternativas, revisar riesgos y explorar escenarios, pero no delego la decisión final. La responsabilidad "
        "del líder es combinar evidencia, experiencia, intuición y conversación con el equipo."
    ),
    "“¿Cuáles son las técnicas utiliza para motivar a su equipo y mantener altos niveles de desempeño?”": (
        "Me funciona conectar el trabajo diario con un propósito claro, dar autonomía real y reconocer avances concretos. Un desarrollador se motiva más cuando entiende el impacto "
        "de lo que construye, tiene espacio para proponer y recibe retroalimentación útil. También procuro que las herramientas, incluida la IA, eliminen fricción: automatizar pruebas, "
        "generar borradores de documentación o apoyar revisiones de código puede liberar tiempo para trabajo de mayor valor. La motivación aumenta cuando la tecnología se siente como "
        "un multiplicador del equipo, no como vigilancia o amenaza."
    ),
    "“¿Cómo aborda la falta de motivación en un miembro del equipo o en todo el grupo?”": (
        "Primero intento entender la causa antes de asumir falta de compromiso. Puede haber sobrecarga, poca claridad, problemas personales, tareas mal asignadas o ausencia de reconocimiento. "
        "Hablo de forma directa y respetuosa, reviso expectativas y busco acuerdos concretos. Si la desmotivación es grupal, reviso el sistema: prioridades cambiantes, exceso de reuniones, "
        "falta de foco o herramientas que no ayudan. Muchas veces mejorar la comunicación y reducir trabajo repetitivo con automatización devuelve energía al equipo."
    ),
    "“¿Cómo aborda los conflictos en su equipo? ¿Qué pasos sigue para garantizar una resolución positiva?”": (
        "Los abordo temprano, con conversación directa y centrada en hechos. Primero separo el problema de las personas; luego escucho las versiones, identifico el desacuerdo real y busco "
        "un criterio compartido para decidir. En equipos técnicos, muchos conflictos nacen de supuestos distintos sobre calidad, tiempos o alcance. Documentar decisiones, acuerdos y razones "
        "ayuda mucho. La IA puede apoyar organizando antecedentes o resumiendo puntos de vista, pero la resolución requiere confianza, respeto y una intervención humana cuidadosa."
    ),
    "“¿Cómo asegura que está escuchando y comprendiendo adecuadamente las preocupaciones o ideas de los miembros de su equipo?”": (
        "Uso preguntas abiertas, confirmo con mis propias palabras lo que entendí y evito responder demasiado rápido. También observo señales no verbales y patrones: si varias personas repiten "
        "una preocupación, probablemente hay un problema de fondo. En mi experiencia, escuchar bien implica crear espacios donde una persona pueda disentir sin sentir que queda marcada. "
        "Las herramientas de IA pueden ayudar a registrar acuerdos y pendientes, pero la escucha activa ocurre cuando el líder está presente y dispuesto a cambiar una decisión si la evidencia del equipo lo justifica."
    ),
    "“El entorno empresarial está en constante cambio. ¿Cómo gestiona usted el cambio y la incertidumbre?”": (
        "Gestiono el cambio dividiéndolo en ciclos cortos: entender el contexto, plantear hipótesis, ejecutar, medir y ajustar. En una empresa propia, no siempre se puede esperar certeza total; "
        "hay que construir aprendizaje rápido sin perder responsabilidad. Procuro explicar al equipo qué sabemos, qué no sabemos y cómo vamos a validar. La IA es valiosa en ese proceso porque acelera "
        "investigación, prototipado, análisis de información y comunicación interna, pero el rumbo lo define la estrategia y el juicio del equipo."
    ),
    "“¿Qué habilidades cree que son más importantes para un líder al implementar cambios organizacionales?”": (
        "Las más importantes son comunicación clara, empatía, pensamiento sistémico y capacidad de ejecución. Implementar cambios no es solo anunciar una nueva herramienta o proceso; es entender cómo "
        "afecta roles, hábitos, métricas y confianza. Cuando se introduce IA, por ejemplo, el líder debe explicar para qué se usará, qué límites tendrá, cómo se protegerá el criterio humano y cómo se medirá "
        "su impacto. Así el cambio se vive como una mejora colectiva y no como una imposición."
    ),
    "“¿Qué importancia le da al\xa0aprendizaje continuo\xa0y al desarrollo personal para un líder?”": (
        "Le doy una importancia central. En tecnología, lo que era una ventaja hace pocos años puede volverse estándar rápidamente. Un líder que deja de aprender empieza a decidir con mapas viejos. "
        "El aprendizaje continuo no es solo tomar cursos; también es escuchar al equipo, revisar errores, estudiar tendencias y experimentar con nuevas herramientas de forma responsable. La IA hace aún más "
        "necesario este hábito, porque cambia la forma de programar, documentar, vender, atender clientes y coordinar equipos."
    ),
    "“¿Cuáles son las formas en que usted sigue desarrollando sus propias habilidades gerenciales?”": (
        "Me desarrollo combinando práctica, retroalimentación y estudio. Reviso cómo salen los proyectos, pido feedback al equipo, analizo decisiones difíciles y busco aprender de otros líderes y fundadores. "
        "También experimento con nuevas tecnologías en contextos controlados antes de llevarlas al equipo. En el caso de la IA, procuro entender tanto sus capacidades como sus límites: privacidad, sesgos, calidad "
        "de resultados y dependencia excesiva. Eso me permite adoptarla con criterio y enseñar al equipo a usarla como herramienta de apoyo."
    ),
    "“Para concluir, ¿cuáles considera que son las tres habilidades gerenciales más importantes que todo líder debería tener para tener éxito?”": (
        "Destacaría tres: comunicación, toma de decisiones y desarrollo de personas. La comunicación alinea; la toma de decisiones convierte la visión en acción; y el desarrollo de personas hace que el equipo no "
        "dependa de una sola figura. En empresas tecnológicas agregaría una capacidad transversal: aprender y adaptar nuevas herramientas, especialmente IA, de manera ética y práctica. Pero esa capacidad funciona "
        "mejor cuando fortalece las tres habilidades anteriores."
    ),
    "“¿Hay alguna otra habilidad o consejo que quisiera compartir sobre lo que hace que un líder sea efectivo en la gestión de equipos?”": (
        "Un consejo importante es no confundir velocidad con liderazgo. La velocidad sin claridad crea retrabajo; la claridad sin ejecución se queda en intención. Un líder efectivo equilibra visión, foco y humanidad. "
        "También debe construir sistemas donde el conocimiento circule: documentación, rituales de seguimiento, decisiones visibles y herramientas que ayuden a colaborar. La IA puede ser una gran aliada en ese sistema "
        "si se usa para amplificar capacidades humanas: mejorar comunicación, detectar bloqueos, reducir carga operativa y permitir que el equipo piense mejor."
    ),
    "“Gracias por compartir su tiempo y experiencia con nosotros. Esta información será muy valiosa para conocer la habilidades gerenciales en el mundo de los negocios ]. ¿Hay algo más que le gustaría añadir sobre este tema?”": (
        "Sí. Añadiría que el liderazgo actual requiere una mentalidad de construcción continua. Las empresas que mejor incorporan tecnología no son necesariamente las que reemplazan más personas, sino las que diseñan "
        "mejores formas de trabajar. Para mí, la inteligencia artificial debe verse como una extensión del equipo: ayuda a comunicar, analizar, documentar y crear con mayor velocidad, pero el sentido, la ética, la confianza "
        "y la responsabilidad siguen dependiendo de las personas."
    ),
}


def set_run_font(run, name="Arial", size=10.5, color=None, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def make_response_after(paragraph, answer):
    response = insert_paragraph_after(paragraph)
    response.style = "Normal"
    response.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    response.paragraph_format.space_before = Pt(2)
    response.paragraph_format.space_after = Pt(8)
    response.paragraph_format.line_spacing = 1.08

    label = response.add_run("Respuesta: ")
    set_run_font(label, color=(31, 78, 121), bold=True)

    body = response.add_run(answer)
    set_run_font(body, color=(0, 0, 0))


def add_context_note(doc):
    title = next((p for p in doc.paragraphs if p.text.strip().startswith("Guion de Entrevista")), None)
    if not title:
        return
    note = insert_paragraph_after(title)
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(12)
    run = note.add_run(
        "Contexto de respuesta: desarrollador, líder de equipo y cofundador de una empresa tecnológica, "
        "con énfasis en adopción responsable de IA para fortalecer equipos, comunicación y toma de decisiones."
    )
    set_run_font(run, size=10, color=(89, 89, 89), italic=True)


def main():
    doc = Document(DOCX_PATH)
    add_context_note(doc)

    inserted = 0
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        answer = ANSWERS.get(text)
        if answer:
            make_response_after(paragraph, answer)
            inserted += 1

    if inserted != len(ANSWERS):
        missing = len(ANSWERS) - inserted
        raise RuntimeError(f"Only inserted {inserted} responses; {missing} answer(s) were not matched.")

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
